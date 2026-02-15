import io
from unittest.mock import MagicMock, patch

import pytest


class TestPDFProcessor:
    def test_extract_text_valid_pdf(self, sample_pdf_bytes):
        from processors.pdf_processor import extract_text

        result = extract_text(sample_pdf_bytes)
        assert "test PDF document" in result
        assert "project deadlines" in result

    def test_extract_text_invalid_bytes(self):
        from processors.pdf_processor import extract_text

        result = extract_text(b"not a pdf")
        assert result == ""

    def test_extract_text_empty_pdf(self):
        import fitz

        doc = fitz.open()
        doc.new_page()  # blank page, no text
        pdf_bytes = doc.tobytes()
        doc.close()

        from processors.pdf_processor import extract_text

        result = extract_text(pdf_bytes)
        assert result == ""


class TestTextProcessor:
    def test_extract_plain_text_utf8(self):
        from processors.text_processor import extract_text

        content = b"Hello, this is a test document."
        result = extract_text(content, "test.txt")
        assert result == "Hello, this is a test document."

    def test_extract_plain_text_latin1(self):
        from processors.text_processor import extract_text

        content = "caf\xe9 r\xe9sum\xe9".encode("latin-1")
        result = extract_text(content, "test.txt")
        assert "caf" in result

    def test_extract_markdown(self):
        from processors.text_processor import extract_text

        content = b"# Title\n\nSome **bold** text"
        result = extract_text(content, "readme.md")
        assert "# Title" in result

    def test_extract_docx(self):
        from docx import Document
        from processors.text_processor import extract_text

        doc = Document()
        doc.add_paragraph("Test paragraph one")
        doc.add_paragraph("Test paragraph two")
        buf = io.BytesIO()
        doc.save(buf)
        result = extract_text(buf.getvalue(), "test.docx")
        assert "Test paragraph one" in result
        assert "Test paragraph two" in result

    def test_extract_email(self):
        from processors.text_processor import extract_text

        eml = b"""From: alice@example.com
To: bob@example.com
Subject: Meeting Notes
Date: Mon, 1 Jan 2025 10:00:00 +0000
Content-Type: text/plain

Here are the meeting notes from today."""
        result = extract_text(eml, "meeting.eml")
        assert "Meeting Notes" in result
        assert "meeting notes from today" in result


class TestCalendarProcessor:
    def test_extract_events(self, sample_ics_bytes):
        from processors.calendar_processor import extract_text

        result = extract_text(sample_ics_bytes)
        assert "Team Meeting" in result
        assert "Room 42" in result
        assert "Weekly sync" in result

    def test_extract_empty_calendar(self):
        from processors.calendar_processor import extract_text

        ics = b"""BEGIN:VCALENDAR
VERSION:2.0
PRODID:-//Test//EN
END:VCALENDAR"""
        result = extract_text(ics)
        assert result == "Empty calendar file"

    def test_extract_invalid_ics(self):
        from processors.calendar_processor import extract_text

        result = extract_text(b"not an ics file")
        # Should return empty string or handle gracefully
        assert isinstance(result, str)


class TestImageProcessor:
    def test_extract_text_with_mocked_captioner(self):
        """Mock the BLIP pipeline to avoid loading the actual model."""
        import sys

        from PIL import Image

        # Ensure transformers module is available (mock it if not installed)
        if "transformers" not in sys.modules:
            sys.modules["transformers"] = MagicMock()
        if "processors.image_processor" in sys.modules:
            del sys.modules["processors.image_processor"]

        from processors import image_processor

        mock_pipeline = MagicMock(return_value=[{"generated_text": "a photo of a cat"}])
        with patch.object(
            image_processor, "_get_captioner", return_value=mock_pipeline
        ):
            img = Image.new("RGB", (1, 1), color="red")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            result = image_processor.extract_text(buf.getvalue())
            assert "cat" in result
            mock_pipeline.assert_called_once()


class TestAudioProcessor:
    def test_extract_text_with_mocked_whisper(self):
        """Mock Whisper to avoid loading the actual model."""
        import sys

        # Ensure whisper module is available (mock it if not installed)
        if "whisper" not in sys.modules:
            sys.modules["whisper"] = MagicMock()
        if "processors.audio_processor" in sys.modules:
            del sys.modules["processors.audio_processor"]

        from processors import audio_processor

        mock_model = MagicMock()
        mock_model.transcribe.return_value = {"text": "Hello from the audio"}
        with patch.object(audio_processor, "_get_model", return_value=mock_model):
            result = audio_processor.extract_text(b"fake audio bytes", "test.mp3")
            assert "Hello from the audio" in result
